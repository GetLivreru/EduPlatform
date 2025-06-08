from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Form
from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId
from typing import List, Optional
import os
from dotenv import load_dotenv
from datetime import datetime
from openai import OpenAI
from ..models import User, UserRole, DocumentS3
from ..middleware import require_teacher_or_admin, get_current_user
from ..s3_service import s3_service
import json
import io
import traceback
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Пробуем разные способы импорта PyMuPDF
try:
    import fitz  # PyMuPDF for PDF processing
    PYMUPDF_AVAILABLE = True
    logger.info("✅ PyMuPDF успешно импортирован")
except ImportError:
    try:
        import pymupdf as fitz  # Новый способ импорта
        PYMUPDF_AVAILABLE = True
        logger.info("✅ PyMuPDF (pymupdf) успешно импортирован")
    except ImportError:
        PYMUPDF_AVAILABLE = False
        logger.warning("❌ PyMuPDF не установлен. PDF файлы не будут поддерживаться.")

try:
    from docx import Document  # python-docx for Word documents
    logger.info("✅ python-docx успешно импортирован")
except ImportError:
    logger.error("❌ python-docx не установлен")

# Load .env from parent directory with encoding fallback
env_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env')
logger.info(f"🔍 Поиск .env файла по пути: {env_path}")

try:
    load_dotenv(env_path, encoding='utf-8')
    logger.info("✅ .env файл загружен (UTF-8)")
except UnicodeDecodeError:
    try:
        load_dotenv(env_path, encoding='utf-16')
        logger.info("✅ .env файл загружен (UTF-16)")
    except Exception as e:
        logger.warning(f"⚠️ Ошибка загрузки .env файла: {e}")
except FileNotFoundError:
    logger.error("❌ .env файл не найден!")

router = APIRouter()

# MongoDB connection
MONGODB_URL = os.getenv("MONGODB_URL", "mongodb+srv://Lida:oayjqe2005@cluster0.ejidejg.mongodb.net/?retryWrites=true&w=majority")

logger.info(f"🗄️ MongoDB URL: {MONGODB_URL}")

client = AsyncIOMotorClient(MONGODB_URL)
db = client.LearnApp

# OpenAI configuration
openai_api_key = os.getenv("OPENAI_API_KEY")
if openai_api_key:
    if openai_api_key == "your_openai_api_key_here":
        logger.error("❌ OPENAI_API_KEY не настроен! Используется placeholder значение")
        openai_client = None
    else:
        openai_client = OpenAI(api_key=openai_api_key)
        logger.info("✅ OpenAI клиент инициализирован")
else:
    logger.error("❌ OPENAI_API_KEY не найден в переменных окружения")
    openai_client = None

async def extract_text_from_file(file: UploadFile) -> str:
    """Извлекает текст из загруженного файла"""
    logger.info(f"📄 Начало обработки файла: {file.filename}, тип: {file.content_type}, размер: {file.size}")
    
    try:
        content = await file.read()
        logger.info(f"📖 Содержимое файла прочитано, размер: {len(content)} байт")
        
        if file.content_type == "application/pdf":
            # Обработка PDF файлов
            logger.info("🔄 Обработка PDF файла...")
            if not PYMUPDF_AVAILABLE:
                logger.error("❌ PyMuPDF недоступен для обработки PDF")
                raise HTTPException(
                    status_code=400, 
                    detail="PDF файлы не поддерживаются. PyMuPDF не установлен. Используйте DOCX или TXT файлы."
                )
            
            pdf_document = fitz.open("pdf", content)
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text()
            pdf_document.close()
            logger.info(f"✅ PDF обработан, извлечено {len(text)} символов")
            return text
            
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Обработка Word документов
            logger.info("🔄 Обработка Word документа...")
            doc = Document(io.BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            logger.info(f"✅ Word документ обработан, извлечено {len(text)} символов")
            return text
            
        elif file.content_type == "text/plain":
            # Обработка текстовых файлов
            logger.info("🔄 Обработка текстового файла...")
            try:
                text = content.decode('utf-8')
                logger.info(f"✅ Текстовый файл обработан (UTF-8), извлечено {len(text)} символов")
                return text
            except UnicodeDecodeError:
                try:
                    text = content.decode('cp1251')  # Кириллица Windows
                    logger.info(f"✅ Текстовый файл обработан (CP1251), извлечено {len(text)} символов")
                    return text
                except UnicodeDecodeError:
                    text = content.decode('latin-1')  # Fallback
                    logger.info(f"✅ Текстовый файл обработан (Latin-1), извлечено {len(text)} символов")
                    return text
            
        else:
            supported_formats = ["TXT", "DOCX"]
            if PYMUPDF_AVAILABLE:
                supported_formats.insert(0, "PDF")
            
            logger.error(f"❌ Неподдерживаемый тип файла: {file.content_type}")
            raise HTTPException(
                status_code=400, 
                detail=f"Неподдерживаемый тип файла. Поддерживаются: {', '.join(supported_formats)}"
            )
            
    except Exception as e:
        logger.error(f"❌ Ошибка при обработке файла: {str(e)}")
        logger.error(f"❌ Трассировка: {traceback.format_exc()}")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке файла: {str(e)}")

async def extract_text_from_bytes(file_content: bytes, content_type: str, filename: str) -> str:
    """Извлекает текст из байтов файла (для S3)"""
    logger.info(f"📄 Начало обработки файла из S3: {filename}, тип: {content_type}, размер: {len(file_content)} байт")
    
    try:
        if content_type == "application/pdf":
            # Обработка PDF файлов
            logger.info("🔄 Обработка PDF файла...")
            if not PYMUPDF_AVAILABLE:
                logger.error("❌ PyMuPDF недоступен для обработки PDF")
                raise HTTPException(
                    status_code=400, 
                    detail="PDF файлы не поддерживаются. PyMuPDF не установлен. Используйте DOCX или TXT файлы."
                )
            
            pdf_document = fitz.open("pdf", file_content)
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text()
            pdf_document.close()
            logger.info(f"✅ PDF обработан, извлечено {len(text)} символов")
            return text
            
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Обработка Word документов
            logger.info("🔄 Обработка Word документа...")
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            logger.info(f"✅ Word документ обработан, извлечено {len(text)} символов")
            return text
            
        elif content_type == "text/plain":
            # Обработка текстовых файлов
            logger.info("🔄 Обработка текстового файла...")
            try:
                text = file_content.decode('utf-8')
                logger.info(f"✅ Текстовый файл обработан (UTF-8), извлечено {len(text)} символов")
                return text
            except UnicodeDecodeError:
                try:
                    text = file_content.decode('cp1251')  # Кириллица Windows
                    logger.info(f"✅ Текстовый файл обработан (CP1251), извлечено {len(text)} символов")
                    return text
                except UnicodeDecodeError:
                    text = file_content.decode('latin-1')  # Fallback
                    logger.info(f"✅ Текстовый файл обработан (Latin-1), извлечено {len(text)} символов")
                    return text
            
        else:
            supported_formats = ["TXT", "DOCX"]
            if PYMUPDF_AVAILABLE:
                supported_formats.insert(0, "PDF")
            
            logger.error(f"❌ Неподдерживаемый тип файла: {content_type}")
            raise HTTPException(
                status_code=400, 
                detail=f"Неподдерживаемый тип файла. Поддерживаются: {', '.join(supported_formats)}"
            )
            
    except Exception as e:
        logger.error(f"❌ Ошибка при обработке файла: {str(e)}")
        logger.error(f"❌ Трассировка: {traceback.format_exc()}")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке файла: {str(e)}")

async def generate_quiz_with_gpt(document_text: str, quiz_title: str, difficulty: str, questions_count: int) -> dict:
    """Генерирует квиз с помощью GPT на основе документа"""
    logger.info(f"🤖 Начало генерации квиза: заголовок='{quiz_title}', сложность={difficulty}, вопросов={questions_count}")
    
    try:
        # Проверяем наличие OpenAI API ключа
        if not openai_api_key or openai_api_key == "your_openai_api_key_here":
            logger.error("❌ OpenAI API ключ не настроен")
            raise HTTPException(
                status_code=400, 
                detail="OpenAI API ключ не настроен. Пожалуйста, добавьте OPENAI_API_KEY в файл .env"
            )
        
        if not openai_client:
            logger.error("❌ OpenAI клиент не инициализирован")
            raise HTTPException(
                status_code=400, 
                detail="OpenAI клиент не инициализирован. Проверьте OPENAI_API_KEY"
            )
        
        # Ограничиваем текст документа для GPT
        max_text_length = 8000  # Ограничение для API
        original_length = len(document_text)
        if len(document_text) > max_text_length:
            document_text = document_text[:max_text_length] + "..."
            logger.info(f"📏 Текст обрезан с {original_length} до {len(document_text)} символов")
        
        prompt = f"""
        На основе следующего документа создай тест из {questions_count} вопросов уровня сложности "{difficulty}".
        
        Документ:
        {document_text}
        
        Создай JSON объект со следующей структурой:
        {{
            "title": "{quiz_title}",
            "description": "Тест создан на основе загруженного документа",
            "category": "Документ",
            "difficulty": "{difficulty}",
            "time_limit": {questions_count * 2},
            "questions": [
                {{
                    "question": "Текст вопроса",
                    "options": ["Вариант 1", "Вариант 2", "Вариант 3", "Вариант 4"],
                    "correct_answer": 0
                }}
            ]
        }}
        
        Требования:
        - Вопросы должны быть основаны на содержании документа
        - У каждого вопроса должно быть 4 варианта ответа
        - correct_answer - это индекс правильного ответа (0-3)
        - Вопросы должны соответствовать уровню сложности {difficulty}
        - Ответь только JSON, без дополнительного текста
        """
        
        logger.info("🌐 Отправка запроса в OpenAI API...")
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Ты помощник преподавателя, который создает образовательные тесты на основе документов."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7
        )
        
        logger.info("✅ Ответ получен от OpenAI API")
        response_content = response.choices[0].message.content
        logger.info(f"📝 Длина ответа: {len(response_content)} символов")
        
        # Очищаем ответ от markdown-форматирования
        cleaned_content = response_content.strip()
        
        # Удаляем маркеры markdown блока кода, если они есть
        if cleaned_content.startswith("```json"):
            cleaned_content = cleaned_content[7:]  # Удаляем "```json"
        elif cleaned_content.startswith("```"):
            cleaned_content = cleaned_content[3:]   # Удаляем "```"
            
        if cleaned_content.endswith("```"):
            cleaned_content = cleaned_content[:-3]  # Удаляем закрывающие "```"
            
        cleaned_content = cleaned_content.strip()
        logger.info(f"🧹 Очищенный JSON (первые 200 символов): {cleaned_content[:200]}...")
        
        try:
            quiz_data = json.loads(cleaned_content)
            logger.info(f"✅ JSON успешно разобран, квиз содержит {len(quiz_data.get('questions', []))} вопросов")
            return quiz_data
        except json.JSONDecodeError as json_err:
            logger.error(f"❌ Ошибка разбора JSON: {json_err}")
            logger.error(f"❌ Очищенный ответ: {cleaned_content}")
            raise HTTPException(
                status_code=500, 
                detail="Ошибка обработки ответа от ИИ. Попробуйте еще раз"
            )
        
    except Exception as e:
        # Обрабатываем различные типы ошибок OpenAI
        error_message = str(e)
        logger.error(f"❌ Ошибка генерации квиза: {error_message}")
        logger.error(f"❌ Трассировка: {traceback.format_exc()}")
        
        if "authentication" in error_message.lower() or "unauthorized" in error_message.lower():
            raise HTTPException(
                status_code=400, 
                detail="Неверный OpenAI API ключ. Проверьте настройки в файле .env"
            )
        elif "rate_limit" in error_message.lower() or "quota" in error_message.lower():
            raise HTTPException(
                status_code=429, 
                detail="Превышен лимит запросов к OpenAI API. Попробуйте позже"
            )
        elif "json" in error_message.lower():
            raise HTTPException(
                status_code=500, 
                detail="Ошибка обработки ответа от ИИ. Попробуйте еще раз"
            )
        else:
            # Если это HTTPException, пробрасываем дальше
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(status_code=500, detail=f"Ошибка при генерации квиза: {str(e)}")

@router.post("/upload-document",
            summary="Загрузить документ и создать квиз [преподаватель]",
            description="Загружает документ и генерирует квиз с помощью ИИ (только для преподавателей)")
async def upload_document_and_generate_quiz(
    file: UploadFile = File(...),
    quiz_title: str = Form(...),
    difficulty: str = Form(...),
    questions_count: int = Form(5),
    current_user: User = Depends(get_current_user)
):
    logger.info(f"🚀 Начало загрузки документа от пользователя: {current_user.name} (роль: {current_user.role})")
    logger.info(f"📋 Параметры: файл={file.filename}, заголовок='{quiz_title}', сложность={difficulty}, вопросов={questions_count}")
    
    # Проверяем, что пользователь - преподаватель (НЕ админ!)
    if not current_user.role or current_user.role not in ['teacher']:
        logger.error(f"❌ Отказано в доступе для пользователя {current_user.name} с ролью {current_user.role}")
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут загружать документы и создавать квизы"
        )
    
    try:
        # Валидация размера файла (максимум 10MB)
        if file.size > 10 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail="Файл слишком большой. Максимальный размер: 10MB"
            )
        
        # Валидация типа файла
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain"
        ]
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail="Неподдерживаемый тип файла. Поддерживаются: PDF, DOCX, TXT"
            )
        
        logger.info("🔄 Извлечение текста из файла...")
        # Извлекаем текст из файла для анализа
        document_text = await extract_text_from_file(file)
        
        if len(document_text.strip()) < 100:
            logger.error(f"❌ Документ слишком короткий: {len(document_text.strip())} символов")
            raise HTTPException(status_code=400, detail="Документ слишком короткий для создания качественного теста")
        
        logger.info(f"✅ Текст извлечен, длина: {len(document_text)} символов")
        
        # Сбрасываем указатель файла для загрузки в S3
        await file.seek(0)
        
        logger.info("☁️ Загрузка файла в AWS S3...")
        # Загружаем файл в S3
        if not s3_service.is_available():
            raise HTTPException(
                status_code=500,
                detail="S3 сервис недоступен. Проверьте настройки AWS."
            )
        
        s3_metadata = await s3_service.upload_file(file, current_user.id)
        logger.info(f"✅ Файл загружен в S3: {s3_metadata['s3_key']}")
        
        logger.info("🗄️ Сохранение метаданных документа в БД...")
        # Сохраняем метаданные документа в MongoDB
        document_info = {
            "s3_key": s3_metadata["s3_key"],
            "s3_bucket": s3_metadata["s3_bucket"],
            "s3_region": s3_metadata["s3_region"],
            "original_filename": s3_metadata["original_filename"],
            "content_type": s3_metadata["content_type"],
            "file_size": s3_metadata["file_size"],
            "uploaded_by": current_user.id,
            "uploaded_at": s3_metadata["uploaded_at"],
            "text_length": len(document_text)
        }
        
        document_result = await db.documents.insert_one(document_info)
        logger.info(f"✅ Метаданные документа сохранены в БД с ID: {document_result.inserted_id}")
        
        logger.info("🤖 Генерация квиза с помощью GPT...")
        # Генерируем квиз с помощью GPT
        quiz_data = await generate_quiz_with_gpt(document_text, quiz_title, difficulty, questions_count)
        
        # Добавляем информацию о создателе и источнике
        quiz_data["created_by"] = current_user.id
        quiz_data["source_document_id"] = str(document_result.inserted_id)
        quiz_data["created_at"] = datetime.utcnow()
        quiz_data["updated_at"] = datetime.utcnow()
        
        logger.info("🗄️ Сохранение квиза в БД...")
        # Сохраняем квиз в базе данных
        quiz_result = await db.quizzes.insert_one(quiz_data)
        quiz_data["_id"] = str(quiz_result.inserted_id)
        
        logger.info(f"✅ Квиз успешно создан с ID: {quiz_result.inserted_id}")
        
        return {
            "message": "Документ успешно загружен в S3 и квиз создан",
            "document_id": str(document_result.inserted_id),
            "quiz_id": str(quiz_result.inserted_id),
            "s3_key": s3_metadata["s3_key"],
            "quiz": quiz_data
        }
        
    except Exception as e:
        logger.error(f"❌ Общая ошибка при загрузке документа: {str(e)}")
        logger.error(f"❌ Трассировка: {traceback.format_exc()}")
        
        # Если ошибка произошла после загрузки в S3, пытаемся удалить файл
        if 's3_metadata' in locals() and s3_metadata.get('s3_key'):
            logger.info(f"🧹 Очистка: удаляем файл {s3_metadata['s3_key']} из S3")
            await s3_service.delete_file(s3_metadata['s3_key'])
        
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/my-documents",
           summary="Получить мои документы [преподаватель]",
           description="Возвращает список документов, загруженных текущим преподавателем")
async def get_my_documents(current_user: User = Depends(get_current_user)):
    # Проверяем, что пользователь - преподаватель
    if not current_user.role or current_user.role not in ['teacher']:
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут просматривать свои документы"
        )
    
    try:
        documents = []
        cursor = db.documents.find({"uploaded_by": current_user.id})
        async for doc in cursor:
            doc["id"] = str(doc["_id"])
            del doc["_id"]
            
            # Находим связанные квизы
            quizzes = await db.quizzes.find({"source_document_id": doc["id"]}).to_list(None)
            doc["generated_quizzes"] = len(quizzes)
            
            # Генерируем временную ссылку для скачивания (если S3 доступен)
            if s3_service.is_available() and doc.get("s3_key"):
                doc["download_url"] = s3_service.get_file_url(doc["s3_key"], expiration=3600)  # 1 час
            else:
                doc["download_url"] = None
            
            # Проверяем существование файла в S3
            if doc.get("s3_key"):
                metadata = await s3_service.get_file_metadata(doc["s3_key"])
                doc["file_exists"] = metadata is not None
            else:
                doc["file_exists"] = False
            
            documents.append(doc)
        
        return {"documents": documents}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/my-generated-quizzes",
           summary="Получить мои сгенерированные квизы [преподаватель]",
           description="Возвращает список квизов, созданных текущим преподавателем")
async def get_my_generated_quizzes(current_user: User = Depends(get_current_user)):
    # Проверяем, что пользователь - преподаватель
    if not current_user.role or current_user.role not in ['teacher']:
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут просматривать свои квизы"
        )
    
    try:
        quizzes = []
        cursor = db.quizzes.find({"created_by": current_user.id})
        async for quiz in cursor:
            quiz["id"] = str(quiz["_id"])
            del quiz["_id"]
            
            # Добавляем статистику по попыткам
            attempts_count = await db.quiz_attempts.count_documents({"quiz_id": quiz["id"]})
            quiz["attempts_count"] = attempts_count
            
            quizzes.append(quiz)
        
        return {"quizzes": quizzes}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/documents/{document_id}",
              summary="Удалить документ [преподаватель]",
              description="Удаляет документ и все связанные с ним квизы")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    # Проверяем, что пользователь - преподаватель
    if not current_user.role or current_user.role not in ['teacher']:
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут удалять свои документы"
        )
    
    try:
        # Проверяем, что документ принадлежит текущему пользователю
        document = await db.documents.find_one({
            "_id": ObjectId(document_id),
            "uploaded_by": current_user.id
        })
        
        if not document:
            raise HTTPException(status_code=404, detail="Документ не найден или вы не являетесь его автором")
        
        # Удаляем файл из S3 (если есть s3_key)
        if document.get("s3_key") and s3_service.is_available():
            logger.info(f"🗑️ Удаление файла {document['s3_key']} из S3...")
            success = await s3_service.delete_file(document["s3_key"])
            if success:
                logger.info(f"✅ Файл {document['s3_key']} удален из S3")
            else:
                logger.warning(f"⚠️ Не удалось удалить файл {document['s3_key']} из S3")
        
        # Удаляем связанные квизы
        deleted_quizzes = await db.quizzes.delete_many({"source_document_id": document_id})
        logger.info(f"🗑️ Удалено {deleted_quizzes.deleted_count} связанных квизов")
        
        # Удаляем метаданные документа из MongoDB
        await db.documents.delete_one({"_id": ObjectId(document_id)})
        logger.info(f"✅ Метаданные документа {document_id} удалены из БД")
        
        return {
            "message": "Документ и связанные квизы успешно удалены",
            "deleted_quizzes": deleted_quizzes.deleted_count,
            "file_deleted_from_s3": document.get("s3_key") is not None
        }
        
    except Exception as e:
        logger.error(f"❌ Ошибка при удалении документа: {str(e)}")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/quiz-stats/{quiz_id}",
           summary="Статистика по квизу [преподаватель]",
           description="Возвращает статистику попыток по квизу")
async def get_quiz_stats(
    quiz_id: str,
    current_user: User = Depends(get_current_user)
):
    # Проверяем, что пользователь - преподаватель
    if not current_user.role or current_user.role not in ['teacher']:
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут просматривать статистику своих квизов"
        )
    
    try:
        # Проверяем, что квиз принадлежит текущему пользователю
        quiz = await db.quizzes.find_one({
            "_id": ObjectId(quiz_id),
            "created_by": current_user.id
        })
        
        if not quiz:
            raise HTTPException(status_code=404, detail="Квиз не найден или вы не являетесь его автором")
        
        # Получаем статистику попыток
        attempts = await db.quiz_attempts.find({"quiz_id": quiz_id}).to_list(None)
        
        total_attempts = len(attempts)
        if total_attempts == 0:
            return {
                "quiz_title": quiz["title"],
                "total_attempts": 0,
                "average_score": 0,
                "completion_rate": 0,
                "attempts": []
            }
        
        completed_attempts = [a for a in attempts if a.get("completed", False)]
        completion_rate = len(completed_attempts) / total_attempts * 100
        
        scores = [a.get("score", 0) for a in completed_attempts if a.get("score") is not None]
        average_score = sum(scores) / len(scores) if scores else 0
        
        return {
            "quiz_title": quiz["title"],
            "total_attempts": total_attempts,
            "completed_attempts": len(completed_attempts),
            "average_score": round(average_score, 2),
            "completion_rate": round(completion_rate, 2),
            "attempts": attempts
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/download/{document_id}",
           summary="Скачать документ [преподаватель]",
           description="Возвращает временную ссылку для скачивания документа из S3")
async def download_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    # Проверяем, что пользователь - преподаватель
    if not current_user.role or current_user.role not in ['teacher']:
        raise HTTPException(
            status_code=403, 
            detail="Только преподаватели могут скачивать свои документы"
        )
    
    try:
        # Проверяем, что документ принадлежит текущему пользователю
        document = await db.documents.find_one({
            "_id": ObjectId(document_id),
            "uploaded_by": current_user.id
        })
        
        if not document:
            raise HTTPException(status_code=404, detail="Документ не найден или вы не являетесь его автором")
        
        # Проверяем наличие S3 ключа
        if not document.get("s3_key"):
            raise HTTPException(status_code=404, detail="Файл не найден в системе хранения")
        
        # Проверяем доступность S3
        if not s3_service.is_available():
            raise HTTPException(status_code=500, detail="Сервис хранения файлов недоступен")
        
        # Генерируем временную ссылку для скачивания (действует 1 час)
        download_url = s3_service.get_file_url(document["s3_key"], expiration=3600)
        
        if not download_url:
            raise HTTPException(status_code=500, detail="Не удалось создать ссылку для скачивания")
        
        return {
            "download_url": download_url,
            "filename": document.get("original_filename", document.get("filename", "document")),
            "expires_in": 3600,  # секунды
            "file_size": document.get("file_size", document.get("size", 0))
        }
        
    except Exception as e:
        logger.error(f"❌ Ошибка при создании ссылки для скачивания: {str(e)}")
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=str(e)) 